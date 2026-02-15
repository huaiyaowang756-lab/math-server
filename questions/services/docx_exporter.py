"""
导出题目为 Word 文档，支持三种模式：
- teacher : 教师模式 — 每题后附答案、分析、详解
- student : 学生模式 — 仅题干
- normal  : 普通模式 — 所有题目在前，答案统一附在最后
"""

import io
import urllib.request
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from lxml import etree

# 答案/解析/详解 灰度底色
_ANSWER_BG_COLOR = "F2F2F2"  # 浅灰

# ──────────────────────── XML 命名空间 ────────────────────────
M_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
MATHML_NS = "http://www.w3.org/1998/Math/MathML"


def _m(tag):
    return f"{{{M_NS}}}{tag}"


def _local(elem):
    tag = elem.tag if isinstance(elem.tag, str) else ""
    return tag.split("}", 1)[1] if "}" in tag else tag


# ──────────────────── MathML → OMML 转换 ────────────────────
def _make_run(text, italic=False):
    """创建 OMML 数学文本节点 <m:r><m:t>text</m:t></m:r>"""
    r = etree.Element(_m("r"))
    if italic:
        rpr = etree.SubElement(r, _m("rPr"))
        sty = etree.SubElement(rpr, _m("sty"))
        sty.set(_m("val"), "i")
    t = etree.SubElement(r, _m("t"))
    t.text = text
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    return r


def _text_content(elem):
    return (elem.text or "").strip()


def _all_text(elem):
    parts = [elem.text or ""]
    for child in elem:
        parts.append(_all_text(child))
        if child.tail:
            parts.append(child.tail)
    return "".join(parts)


def _convert_children(elem, parent):
    for child in elem:
        _convert_mml(child, parent)
        if child.tail and child.tail.strip():
            parent.append(_make_run(child.tail.strip()))


def _convert_mml(elem, parent):
    """递归将 MathML 元素转为 OMML 并追加到 parent"""
    tag = _local(elem)

    # ── 文本节点 ──
    if tag in ("mi", "mn", "mo", "mtext", "ms"):
        text = _text_content(elem)
        if text:
            italic = tag == "mi" and len(text) == 1
            parent.append(_make_run(text, italic=italic))
        # 处理子元素（罕见情况）
        for child in elem:
            _convert_mml(child, parent)
            if child.tail and child.tail.strip():
                parent.append(_make_run(child.tail.strip()))
        return

    # ── 透传容器 ──
    if tag in ("mrow", "mstyle", "mpadded", "merror", "semantics", "annotation-xml"):
        _convert_children(elem, parent)
        return

    # ── 分数 ──
    if tag == "mfrac":
        children = list(elem)
        if len(children) >= 2:
            f = etree.SubElement(parent, _m("f"))
            etree.SubElement(f, _m("fPr"))
            num = etree.SubElement(f, _m("num"))
            den = etree.SubElement(f, _m("den"))
            _convert_mml(children[0], num)
            _convert_mml(children[1], den)
        return

    # ── 上标 ──
    if tag == "msup":
        children = list(elem)
        if len(children) >= 2:
            s = etree.SubElement(parent, _m("sSup"))
            etree.SubElement(s, _m("sSupPr"))
            e = etree.SubElement(s, _m("e"))
            sup = etree.SubElement(s, _m("sup"))
            _convert_mml(children[0], e)
            _convert_mml(children[1], sup)
        return

    # ── 下标 ──
    if tag == "msub":
        children = list(elem)
        if len(children) >= 2:
            s = etree.SubElement(parent, _m("sSub"))
            etree.SubElement(s, _m("sSubPr"))
            e = etree.SubElement(s, _m("e"))
            sub = etree.SubElement(s, _m("sub"))
            _convert_mml(children[0], e)
            _convert_mml(children[1], sub)
        return

    # ── 上下标 ──
    if tag == "msubsup":
        children = list(elem)
        if len(children) >= 3:
            s = etree.SubElement(parent, _m("sSubSup"))
            etree.SubElement(s, _m("sSubSupPr"))
            e = etree.SubElement(s, _m("e"))
            sub = etree.SubElement(s, _m("sub"))
            sup = etree.SubElement(s, _m("sup"))
            _convert_mml(children[0], e)
            _convert_mml(children[1], sub)
            _convert_mml(children[2], sup)
        return

    # ── 平方根 ──
    if tag == "msqrt":
        rad = etree.SubElement(parent, _m("rad"))
        rpr = etree.SubElement(rad, _m("radPr"))
        dh = etree.SubElement(rpr, _m("degHide"))
        dh.set(_m("val"), "1")
        etree.SubElement(rad, _m("deg"))
        e = etree.SubElement(rad, _m("e"))
        _convert_children(elem, e)
        return

    # ── n 次根号 ──
    if tag == "mroot":
        children = list(elem)
        if len(children) >= 2:
            rad = etree.SubElement(parent, _m("rad"))
            etree.SubElement(rad, _m("radPr"))
            deg = etree.SubElement(rad, _m("deg"))
            e = etree.SubElement(rad, _m("e"))
            _convert_mml(children[0], e)
            _convert_mml(children[1], deg)
        return

    # ── 上装饰 (hat, bar, vec 等) ──
    if tag == "mover":
        children = list(elem)
        if len(children) >= 2:
            acc = etree.SubElement(parent, _m("acc"))
            accpr = etree.SubElement(acc, _m("accPr"))
            ch = etree.SubElement(accpr, _m("chr"))
            accent_char = _text_content(children[1])
            ch.set(_m("val"), accent_char or "\u0302")
            e = etree.SubElement(acc, _m("e"))
            _convert_mml(children[0], e)
        return

    # ── 下装饰 ──
    if tag == "munder":
        children = list(elem)
        if len(children) >= 2:
            lim = etree.SubElement(parent, _m("limLow"))
            etree.SubElement(lim, _m("limLowPr"))
            e = etree.SubElement(lim, _m("e"))
            li = etree.SubElement(lim, _m("lim"))
            _convert_mml(children[0], e)
            _convert_mml(children[1], li)
        return

    # ── 上下组合 (∑ ∫ 等带上下限) ──
    if tag == "munderover":
        children = list(elem)
        if len(children) >= 3:
            base_text = _all_text(children[0])
            if base_text.strip() in ("∑", "∫", "∏", "⋃", "⋂", "∮"):
                nary = etree.SubElement(parent, _m("nary"))
                narypr = etree.SubElement(nary, _m("naryPr"))
                c = etree.SubElement(narypr, _m("chr"))
                c.set(_m("val"), base_text.strip())
                sub = etree.SubElement(nary, _m("sub"))
                sup = etree.SubElement(nary, _m("sup"))
                e = etree.SubElement(nary, _m("e"))
                _convert_mml(children[1], sub)
                _convert_mml(children[2], sup)
            else:
                lim = etree.SubElement(parent, _m("limLow"))
                etree.SubElement(lim, _m("limLowPr"))
                e = etree.SubElement(lim, _m("e"))
                li = etree.SubElement(lim, _m("lim"))
                _convert_mml(children[0], e)
                _convert_mml(children[1], li)
                # 上限作为上标
                ssup = etree.SubElement(parent, _m("sSup"))
                etree.SubElement(ssup, _m("sSupPr"))
                e2 = etree.SubElement(ssup, _m("e"))
                sup2 = etree.SubElement(ssup, _m("sup"))
                _convert_mml(children[2], sup2)
        return

    # ── 括号 / 围栏 ──
    if tag == "mfenced":
        open_ch = elem.get("open", "(")
        close_ch = elem.get("close", ")")
        d = etree.SubElement(parent, _m("d"))
        dpr = etree.SubElement(d, _m("dPr"))
        beg = etree.SubElement(dpr, _m("begChr"))
        beg.set(_m("val"), open_ch)
        end = etree.SubElement(dpr, _m("endChr"))
        end.set(_m("val"), close_ch)
        e = etree.SubElement(d, _m("e"))
        _convert_children(elem, e)
        return

    # ── 矩阵 / 表格 ──
    if tag == "mtable":
        mm = etree.SubElement(parent, _m("m"))
        etree.SubElement(mm, _m("mPr"))
        for row_elem in elem:
            if _local(row_elem) in ("mtr", "mlabeledtr"):
                mr = etree.SubElement(mm, _m("mr"))
                for cell in row_elem:
                    if _local(cell) == "mtd":
                        e = etree.SubElement(mr, _m("e"))
                        _convert_children(cell, e)
        return

    # ── 忽略元素 ──
    if tag in ("mspace", "mphantom", "annotation", "none"):
        return

    # ── 回退：提取所有文本 ──
    text = _all_text(elem).strip()
    if text:
        parent.append(_make_run(text))


def _latex_to_omml(latex_str):
    """将 LaTeX 字符串转为 OMML etree Element，失败返回 None"""
    try:
        import latex2mathml.converter
        mathml_str = latex2mathml.converter.convert(latex_str)
        mathml_tree = etree.fromstring(mathml_str.encode("utf-8"))
        omath = etree.Element(_m("oMath"))
        tag = _local(mathml_tree)
        if tag == "math":
            _convert_children(mathml_tree, omath)
        else:
            _convert_mml(mathml_tree, omath)
        return omath
    except Exception:
        return None


# ──────────────────── 字体工具 ────────────────────

_FONT_NAME = "方正仿宋"


def _set_font(run_or_style, name=_FONT_NAME):
    """
    同时设置 ASCII (w:ascii/w:hAnsi) 和 East Asian (w:eastAsia) 字体。
    python-docx 的 font.name 只设置前者，中文走 eastAsia 需要手动补。
    """
    from docx.oxml.ns import qn
    run_or_style.font.name = name
    # 找到 rPr 下的 rFonts 元素，补上 eastAsia
    rpr = getattr(run_or_style, "_element", None)
    if rpr is None:
        return
    # run._element 是 <w:r>, style._element 是 <w:style>
    rFonts = rpr.find(f".//{{{W_NS}}}rFonts")
    if rFonts is None:
        # 先通过 font.name setter 确保 rFonts 已创建
        rFonts = rpr.find(f".//{{{W_NS}}}rFonts")
    if rFonts is not None:
        rFonts.set(qn("w:eastAsia"), name)


# ──────────────────── 内容块写入 Word ────────────────────

_QUESTION_TYPE_MAP = {
    "single_choice": "单选题",
    "multiple_choice": "多选题",
    "fill_blank": "填空题",
    "solution": "解答题",
}

# 图片宽度上限 16 cm，默认 4 cm
_IMG_MAX_WIDTH_CM = 16
_IMG_DEFAULT_WIDTH_CM = 4
_PX_TO_EMU = 9525  # 1px = 9525 EMU (96 DPI)


def _download_image(url, asset_base_url=""):
    """下载图片，返回 BytesIO 流；失败返回 None"""
    try:
        if not url.startswith(("http://", "https://")):
            base = (asset_base_url or "").rstrip("/")
            url = f"{base}/{url.lstrip('/')}" if base else url
        if not url.startswith(("http://", "https://")):
            return None
        req = urllib.request.Request(url, headers={"User-Agent": "Mozilla/5.0"})
        with urllib.request.urlopen(req, timeout=15) as resp:
            return io.BytesIO(resp.read())
    except Exception:
        return None


def _calc_image_width(block):
    """根据 block 的 width 计算 EMU 宽度"""
    w = block.get("width")
    max_w = Cm(_IMG_MAX_WIDTH_CM)
    if isinstance(w, (int, float)) and w > 0:
        emu = Emu(int(w * _PX_TO_EMU))
        return min(emu, max_w)
    return Cm(_IMG_DEFAULT_WIDTH_CM)


_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_BODY_FONT_SIZE_PT = 12  # 正文默认字号


def _vcenter_image_run(run_elem, image_height_pt, font_size_pt=_BODY_FONT_SIZE_PT):
    """
    调整包含图片的 run 的垂直位置，使图片与文本垂直居中对齐。
    默认情况下内联图片底部在文本基线上，大图片会高出文本很多。
    通过设置 <w:position> 将图片下移，使图片中心 ≈ 文本中心。
    """
    # 文本中心 ≈ 基线 + 字号 * 0.35
    # 图片中心 = 基线 + 图片高度 / 2
    # 需要把图片下移: offset = 图片高度/2 - 字号*0.35
    offset_pt = image_height_pt / 2 - font_size_pt * 0.35
    if offset_pt < 1:
        return  # 图片较小，无需调整

    half_points = -round(offset_pt * 2)  # 负值 = 下移
    w = _W_NS
    rpr = run_elem.find(f"{{{w}}}rPr")
    if rpr is None:
        rpr = etree.SubElement(run_elem, f"{{{w}}}rPr")
        run_elem.insert(0, rpr)
    pos = etree.SubElement(rpr, f"{{{w}}}position")
    pos.set(f"{{{w}}}val", str(half_points))


def _add_inline_image(para, doc, image_stream, width):
    """
    将图片以内联方式插入到现有段落中（不新建段落），并垂直居中对齐。
    利用 doc.add_picture 生成图片 run，然后移入目标段落。
    """
    image_stream.seek(0)
    doc.add_picture(image_stream, width=width)
    # add_picture 在文档末尾创建了一个只含图片的临时段落
    temp_para = doc.paragraphs[-1]

    # 从 <wp:extent cy="..."> 读取实际渲染高度（EMU → pt）
    image_height_pt = None
    for elem in temp_para._element.iter():
        tag = elem.tag if isinstance(elem.tag, str) else ""
        if tag.endswith("}extent"):
            cy = elem.get("cy")
            if cy:
                image_height_pt = int(cy) / 12700  # 1pt = 12700 EMU
            break

    # 将图片 run 移到目标段落，同时调整垂直居中
    for child in list(temp_para._element):
        if child.tag.endswith("}pPr"):
            continue  # 跳过段落属性
        if image_height_pt and child.tag.endswith("}r"):
            _vcenter_image_run(child, image_height_pt)
        para._element.append(child)
    # 删除空的临时段落
    temp_para._element.getparent().remove(temp_para._element)


def _add_blocks_to_paragraph(para, blocks, asset_base_url=""):
    """
    将内容块追加到给定段落中，严格按协议决定是否换行：
    - text 块中的 \\n → 换行（新段落）
    - text / latex / image 块 → 同一段落（内联）
    返回涉及的段落列表（用于统一设置底色等）。
    """
    doc = para.part.document
    paras = [para]

    for block in blocks or []:
        btype = block.get("type", "")

        if btype == "text":
            content = block.get("content", "")
            if not content:
                continue
            # 只在协议明确 \n 时换行
            lines = content.split("\n")
            for j, line in enumerate(lines):
                if j > 0:
                    para = doc.add_paragraph()
                    paras.append(para)
                if line:
                    para.add_run(line)

        elif btype == "latex":
            content = block.get("content", "")
            if not content:
                continue
            omml = _latex_to_omml(content)
            if omml is not None:
                para._element.append(omml)
            else:
                run = para.add_run(f" {content} ")
                run.italic = True
                run.font.name = "Cambria Math"

        elif btype in ("image", "svg"):
            url = (block.get("url") or "").strip()
            if not url:
                continue
            stream = _download_image(url, asset_base_url)
            if stream:
                width = _calc_image_width(block)
                try:
                    _add_inline_image(para, doc, stream, width)
                except Exception:
                    para.add_run("[图片]")
            else:
                para.add_run("[图片]")

    return paras


def _apply_shading(paragraph, fill=_ANSWER_BG_COLOR):
    """给段落设置灰度底色（通过 OOXML）"""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    pPr = paragraph._p.get_or_add_pPr()
    shd = pPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        pPr.append(shd)
    shd.set(qn("w:fill"), fill)


def _write_section(doc, blocks, asset_base_url, label=None, label_color=None, apply_shading=False):
    """写一段带可选标签前缀的内容，可选灰度底色"""
    para = doc.add_paragraph()
    if label:
        run = para.add_run(label)
        run.bold = True
        _set_font(run)
        if label_color:
            run.font.color.rgb = label_color
    paras = _add_blocks_to_paragraph(para, blocks, asset_base_url)
    if apply_shading:
        for p in paras:
            _apply_shading(p)


def _write_question_header(doc, num, q_type_str):
    """写题号行"""
    para = doc.add_paragraph()
    run = para.add_run(f"{num}. ")
    run.bold = True
    _set_font(run)
    if q_type_str:
        run = para.add_run(f"（{q_type_str}）")
        _set_font(run)
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        run.font.size = Pt(10)
    return para


# ──────────────────── 三种导出模式 ────────────────────

def _export_teacher(doc, questions):
    """教师模式：每道题后面紧跟答案、分析、详解"""
    for i, q in enumerate(questions, 1):
        qtype = _QUESTION_TYPE_MAP.get(q.get("questionType", ""), "")
        base = q.get("assetBaseUrl", "")

        # 题号与题干同行
        header = _write_question_header(doc, i, qtype)
        _add_blocks_to_paragraph(header, q.get("questionBody", []), base)

        if q.get("answer"):
            _write_section(doc, q["answer"], base, label="【答案】", apply_shading=True)
        if q.get("analysis"):
            _write_section(doc, q["analysis"], base, label="【分析】", apply_shading=True)
        if q.get("detailedSolution"):
            _write_section(doc, q["detailedSolution"], base, label="【详解】", apply_shading=True)

        doc.add_paragraph()  # 题间空行


def _export_student(doc, questions):
    """学生模式：仅题干"""
    for i, q in enumerate(questions, 1):
        qtype = _QUESTION_TYPE_MAP.get(q.get("questionType", ""), "")
        base = q.get("assetBaseUrl", "")

        # 题号与题干同行
        header = _write_question_header(doc, i, qtype)
        _add_blocks_to_paragraph(header, q.get("questionBody", []), base)
        doc.add_paragraph()


def _export_normal(doc, questions):
    """普通模式：所有题目在前，参考答案统一附在最后"""
    for i, q in enumerate(questions, 1):
        qtype = _QUESTION_TYPE_MAP.get(q.get("questionType", ""), "")
        base = q.get("assetBaseUrl", "")

        # 题号与题干同行
        header = _write_question_header(doc, i, qtype)
        _add_blocks_to_paragraph(header, q.get("questionBody", []), base)
        doc.add_paragraph()

    # 分页 + 答案
    doc.add_page_break()
    h = doc.add_heading("参考答案", level=2)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in h.runs:
        _set_font(run)

    for i, q in enumerate(questions, 1):
        base = q.get("assetBaseUrl", "")
        if q.get("answer"):
            _write_section(doc, q["answer"], base, label=f"{i}. ", apply_shading=True)
        if q.get("analysis"):
            _write_section(doc, q["analysis"], base, label="   【分析】", apply_shading=True)
        if q.get("detailedSolution"):
            _write_section(doc, q["detailedSolution"], base, label="   【详解】", apply_shading=True)
        doc.add_paragraph()


# ──────────────────── 对外接口 ────────────────────

def export_questions_docx(questions, mode="teacher"):
    """
    生成 Word 文档并返回 BytesIO 流。

    Args:
        questions: 题目 dict 列表（to_dict 格式）
        mode: 'teacher' | 'student' | 'normal'

    Returns:
        io.BytesIO — 可直接写入 HTTP 响应
    """
    doc = Document()

    # 默认字体（同时设置 ascii + eastAsia）
    style = doc.styles["Normal"]
    _set_font(style)
    style.font.size = Pt(12)

    # 标题
    title = doc.add_heading("数学试卷", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        _set_font(run)

    mode_sub = {"teacher": "（教师版 — 含答案解析）",
                "student": "（学生版）",
                "normal":  "（答案附后）"}
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run(mode_sub.get(mode, ""))
    _set_font(run)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # 分隔线
    doc.add_paragraph("─" * 40)

    if mode == "teacher":
        _export_teacher(doc, questions)
    elif mode == "student":
        _export_student(doc, questions)
    else:
        _export_normal(doc, questions)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
